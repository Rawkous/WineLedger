"""
Regenerate the winery proposal as Word .docx and a simple .html (Word / Google Docs friendly).

Usage (from repo root):
  python -m pip install -r doc/requirements-docx.txt
  python doc/build_proposal_docx.py
"""
from __future__ import annotations

import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

MD_PATH = Path(__file__).resolve().parent / "proposal-wineledger-winery-2026.md"
OUT_DOCX = Path(__file__).resolve().parent / "proposal-wineledger-winery-2026.docx"
OUT_HTML = Path(__file__).resolve().parent / "proposal-wineledger-winery-2026.html"

SPLIT_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline_runs(paragraph: Paragraph, text: str) -> None:
    if not text:
        return
    for part in SPLIT_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2 and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_bullet_paragraph(doc: Document, line: str) -> None:
    """Use a manual bullet + indent so Word/Google Docs never depend on the 'List Bullet' style."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.5)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_after = Pt(6)
    p.add_run("•  ")
    add_inline_runs(p, line)


def _set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    try:
        doc.core_properties.title = "WineLedger — Winery proposal"
        doc.core_properties.author = "Axel"
    except Exception:
        pass


def build_docx() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    _set_doc_defaults(doc)

    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            continue
        if line.startswith("# "):
            t = line[2:].strip()
            # Use Heading 1 for the main title — works across Word builds; avoid Title-style edge cases.
            h = doc.add_heading(t, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("- "):
            add_bullet_paragraph(doc, line[2:].strip())
            continue
        if line.startswith("*") and line.endswith("*") and line.count("*") == 2:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            r = p.add_run(line[1:-1])
            r.italic = True
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        p.paragraph_format.space_after = Pt(6)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def _html_inline(line: str) -> str:
    out: list[str] = []
    for part in SPLIT_RE.split(line):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            out.append("<strong>" + html.escape(part[2:-2]) + "</strong>")
        elif part.startswith("*") and part.endswith("*") and len(part) > 2 and not part.startswith("**"):
            out.append("<em>" + html.escape(part[1:-1]) + "</em>")
        else:
            out.append(html.escape(part))
    return "".join(out)


def build_html() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    parts: list[str] = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="utf-8">',
        "<title>WineLedger — Winery proposal</title>",
        "<style>",
        "body{font-family:Calibri,Arial,sans-serif;font-size:11pt;max-width:40em;margin:1in auto;line-height:1.35;}",
        "h1{font-size:18pt;margin:0 0 0.6em 0;}",
        "h2{font-size:14pt;margin:1em 0 0.4em 0;}",
        "p{margin:0 0 0.6em 0;}",
        ".bullet{margin:0 0 0.5em 2em;text-indent:-1em;}",
        ".hr{height:1px;border:0;background:#ccc;margin:1em 0;}",
        ".note{margin-top:1.2em;font-style:italic;color:#333;}",
        "</style>",
        "</head>",
        "<body>",
    ]
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line == "---":
            parts.append('<div class="hr"></div>')
            continue
        if line.startswith("# "):
            t = html.escape(line[2:].strip())
            parts.append(f"<h1>{t}</h1>")
            continue
        if line.startswith("## "):
            t = _html_inline(line[3:].strip())
            parts.append(f"<h2>{t}</h2>")
            continue
        if line.startswith("- "):
            inner = _html_inline(line[2:].strip())
            parts.append(f'<p class="bullet">•&nbsp;{inner}</p>')
            continue
        if line.startswith("*") and line.endswith("*") and line.count("*") == 2:
            parts.append(f'<p class="note">{_html_inline(line[1:-1])}</p>')
            continue
        parts.append(f"<p>{_html_inline(line)}</p>")

    parts.append("</body></html>")
    OUT_HTML.write_text("\n".join(parts), encoding="utf-8")


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing: {MD_PATH}", file=sys.stderr)
        return 1
    build_docx()
    build_html()
    print(f"Wrote: {OUT_DOCX}")
    print(f"Wrote: {OUT_HTML}  (open in Word, or upload to Google Drive and open with Google Docs)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
